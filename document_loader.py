"""
Document Loader
文档加载器

This module handles loading and processing various document formats.
本模块处理加载和处理各种文档格式。
"""

from typing import List, Optional
from pathlib import Path
import re

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

from config import CHUNK_SIZE, CHUNK_OVERLAP
from vector_store import Document
from logger_config import setup_logger, log_error

logger = setup_logger(__name__)

# ============================================================================
# Text Loading Functions / 文本加载函数
# ============================================================================

def load_txt(file_path: str, encoding: str = "utf-8") -> str:
    """
    Load text from .txt file
    从.txt文件加载文本
    
    Args:
        file_path: Path to text file / 文本文件路径
        encoding: File encoding / 文件编码
        
    Returns:
        File content as string / 字符串形式的文件内容
    """
    try:
        with open(file_path, "r", encoding=encoding) as f:
            content = f.read()
        
        logger.info(f"Loaded {len(content)} characters from {file_path}")
        return content
    
    except UnicodeDecodeError:
        # Try different encodings / 尝试不同的编码
        for enc in ["gbk", "latin-1", "cp1252"]:
            try:
                with open(file_path, "r", encoding=enc) as f:
                    content = f.read()
                logger.warning(f"File loaded with {enc} encoding")
                return content
            except:
                continue
        
        logger.error(f"Could not decode file: {file_path}")
        raise
    
    except Exception as e:
        log_error(logger, f"load_txt({file_path})", e)
        raise

def load_pdf(file_path: str) -> str:
    """
    Load text from PDF file
    从PDF文件加载文本
    
    Args:
        file_path: Path to PDF file / PDF文件路径
        
    Returns:
        Extracted text / 提取的文本
    """
    if not PYPDF_AVAILABLE:
        raise ImportError("pypdf is required to load PDF files. Install with: pip install pypdf")
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        content = "\n\n".join(text_parts)
        logger.info(f"Loaded {len(content)} characters from {len(reader.pages)} pages in {file_path}")
        
        return content
    
    except Exception as e:
        log_error(logger, f"load_pdf({file_path})", e)
        raise

def load_docx(file_path: str) -> str:
    """
    Load text from DOCX file
    从DOCX文件加载文本
    
    Args:
        file_path: Path to DOCX file / DOCX文件路径
        
    Returns:
        Extracted text / 提取的文本
    """
    if not PYTHON_DOCX_AVAILABLE:
        raise ImportError("python-docx is required to load DOCX files. Install with: pip install python-docx")
    
    try:
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content = "\n\n".join(paragraphs)
        
        logger.info(f"Loaded {len(content)} characters from {len(paragraphs)} paragraphs in {file_path}")
        
        return content
    
    except Exception as e:
        log_error(logger, f"load_docx({file_path})", e)
        raise

def load_document(file_path: str) -> str:
    """
    Load document based on file extension
    根据文件扩展名加载文档
    
    Args:
        file_path: Path to document file / 文档文件路径
        
    Returns:
        Document content / 文档内容
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    ext = path.suffix.lower()
    
    if ext == ".txt":
        return load_txt(file_path)
    elif ext == ".pdf":
        return load_pdf(file_path)
    elif ext == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .txt, .pdf, .docx")

# ============================================================================
# Text Chunking Functions / 文本分块函数
# ============================================================================

def clean_text(text: str) -> str:
    """
    Clean and normalize text
    清理和规范化文本
    
    Args:
        text: Input text / 输入文本
        
    Returns:
        Cleaned text / 清理后的文本
    """
    # Remove excessive whitespace / 移除多余的空白
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters (optional) / 移除特殊字符（可选）
    # text = re.sub(r'[^\w\s.,!?;:()\[\]{}"\'`-]', '', text)
    
    return text.strip()

def chunk_text(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    separators: Optional[List[str]] = None
) -> List[str]:
    """
    Split text into overlapping chunks
    将文本分割成重叠的块
    
    Args:
        text: Input text / 输入文本
        chunk_size: Maximum chunk size in characters / 最大块大小（字符）
        chunk_overlap: Overlap between chunks / 块之间的重叠
        separators: List of separators to split on / 用于分割的分隔符列表
        
    Returns:
        List of text chunks / 文本块列表
    """
    if separators is None:
        # Default separators: paragraph, sentence, word / 默认分隔符：段落、句子、单词
        separators = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]
    
    # Clean text first / 首先清理文本
    text = clean_text(text)
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    
    def split_with_separator(text: str, sep: str) -> List[str]:
        """Split text by separator while keeping separator"""
        if sep not in text:
            return [text]
        parts = text.split(sep)
        # Add separator back except for the last part / 除最后一部分外添加分隔符
        return [part + sep for part in parts[:-1]] + [parts[-1]]
    
    def merge_chunks(parts: List[str]) -> List[str]:
        """Merge parts into chunks respecting size and overlap"""
        result = []
        current_chunk = ""
        
        for part in parts:
            if len(current_chunk) + len(part) <= chunk_size:
                current_chunk += part
            else:
                if current_chunk:
                    result.append(current_chunk)
                    # Start new chunk with overlap / 使用重叠开始新块
                    if chunk_overlap > 0:
                        # Take last chunk_overlap characters / 取最后chunk_overlap个字符
                        overlap_text = current_chunk[-chunk_overlap:]
                        current_chunk = overlap_text + part
                    else:
                        current_chunk = part
                else:
                    # Part is larger than chunk_size, force split / 部分大于chunk_size，强制分割
                    result.append(part[:chunk_size])
                    current_chunk = part[chunk_size:]
        
        if current_chunk:
            result.append(current_chunk)
        
        return result
    
    # Try splitting with each separator / 尝试使用每个分隔符分割
    parts = [text]
    for separator in separators:
        new_parts = []
        for part in parts:
            new_parts.extend(split_with_separator(part, separator))
        parts = new_parts
        
        # Check if we can create good chunks / 检查是否可以创建良好的块
        merged = merge_chunks(parts)
        if all(len(chunk) <= chunk_size for chunk in merged):
            chunks = merged
            break
    
    # If we still don't have good chunks, force split / 如果仍然没有良好的块，强制分割
    if not chunks:
        chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size - chunk_overlap)]
    
    logger.info(f"Split text into {len(chunks)} chunks")
    return chunks

def load_and_chunk_document(
    file_path: str,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    metadata: Optional[dict] = None
) -> List[Document]:
    """
    Load document and split into chunks
    加载文档并分割成块
    
    Args:
        file_path: Path to document / 文档路径
        chunk_size: Maximum chunk size / 最大块大小
        chunk_overlap: Overlap between chunks / 块之间的重叠
        metadata: Additional metadata for all chunks / 所有块的附加元数据
        
    Returns:
        List of Document objects / Document对象列表
    """
    try:
        # Load document / 加载文档
        content = load_document(file_path)
        
        # Split into chunks / 分割成块
        chunks = chunk_text(content, chunk_size, chunk_overlap)
        
        # Create Document objects / 创建Document对象
        path = Path(file_path)
        base_metadata = metadata or {}
        base_metadata.update({
            "source": path.name,
            "file_path": str(path),
            "total_chunks": len(chunks)
        })
        
        documents = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = base_metadata.copy()
            chunk_metadata["chunk_index"] = i
            
            doc = Document(
                content=chunk,
                metadata=chunk_metadata
            )
            documents.append(doc)
        
        logger.info(f"Created {len(documents)} document chunks from {file_path}")
        return documents
    
    except Exception as e:
        log_error(logger, f"load_and_chunk_document({file_path})", e)
        raise

# ============================================================================
# Batch Processing / 批量处理
# ============================================================================

def load_directory(
    directory_path: str,
    file_extensions: Optional[List[str]] = None,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP
) -> List[Document]:
    """
    Load all documents from a directory
    从目录加载所有文档
    
    Args:
        directory_path: Path to directory / 目录路径
        file_extensions: List of extensions to include / 要包含的扩展名列表
        chunk_size: Maximum chunk size / 最大块大小
        chunk_overlap: Overlap between chunks / 块之间的重叠
        
    Returns:
        List of all Document objects / 所有Document对象的列表
    """
    if file_extensions is None:
        file_extensions = [".txt", ".pdf", ".docx"]
    
    path = Path(directory_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Directory not found: {directory_path}")
    
    all_documents = []
    
    for ext in file_extensions:
        files = list(path.glob(f"*{ext}"))
        logger.info(f"Found {len(files)} {ext} files")
        
        for file_path in files:
            try:
                docs = load_and_chunk_document(
                    str(file_path),
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Failed to load {file_path}: {e}")
                continue
    
    logger.info(f"Loaded total of {len(all_documents)} document chunks from {directory_path}")
    return all_documents

# ============================================================================
# Example Usage / 使用示例
# ============================================================================

if __name__ == "__main__":
    # Test document loading / 测试文档加载
    test_text = """This is a test document. It contains multiple sentences.
    
    This is a new paragraph. We want to test the chunking functionality.
    
    And here's another paragraph with some more content to ensure we have enough text for chunking."""
    
    # Test chunking / 测试分块
    chunks = chunk_text(test_text, chunk_size=50, chunk_overlap=10)
    print(f"Created {len(chunks)} chunks:\n")
    for i, chunk in enumerate(chunks):
        print(f"Chunk {i+1}: {chunk}\n")
